"""Extract text content from various document formats."""

import pypdf
from docx import Document


def extract_text_from_pdf(file_stream) -> str:
    """Extract text from a PDF file stream."""
    try:
        pdf_reader = pypdf.PdfReader(file_stream)
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{text.strip()}")
        
        return "\n\n".join(text_parts) if text_parts else "[PDF content could not be extracted]"
    except Exception as e:
        return f"[Error extracting PDF: {str(e)}]"


def extract_text_from_docx(file_stream) -> str:
    """Extract text from a Word DOCX file stream."""
    try:
        doc = Document(file_stream)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        
        return "\n\n".join(paragraphs) if paragraphs else "[Document appears to be empty]"
    except Exception as e:
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_txt(file_stream) -> str:
    """Extract text from a plain text file stream."""
    try:
        content = file_stream.read()
        if isinstance(content, bytes):
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return "[Error: Unable to decode text file]"
        return content
    except Exception as e:
        return f"[Error reading text file: {str(e)}]"


def extract_document_content(file_stream, filename: str) -> str:
    """
    Extract text content from a document based on its file extension.
    
    Args:
        file_stream: File-like object (from request.files)
        filename: Original filename with extension
        
    Returns:
        Extracted text content as string
    """
    file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    
    file_stream.seek(0)
    
    if file_ext in ['txt', 'md']:
        return extract_text_from_txt(file_stream)
    elif file_ext == 'pdf':
        return extract_text_from_pdf(file_stream)
    elif file_ext in ['docx']:
        return extract_text_from_docx(file_stream)
    elif file_ext == 'doc':
        return "[Legacy .doc format not supported. Please save as .docx or .pdf]"
    else:
        return f"[Unsupported file type: {file_ext}]"
